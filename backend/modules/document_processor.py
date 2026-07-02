import io
import re
import logging
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
import docx  # python-docx
from dataclasses import dataclass
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


@dataclass
class TextChunk:
    text: str
    page: int
    chunk_id: int


def _clean_text(text: str) -> str:
    """Remove null bytes, duplicate newlines/spaces, and keep text clean."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def extract_text_from_pdf(pdf_bytes: bytes) -> list[dict]:
    """
    Extract text from each page of a PDF.
    Returns list of {"page": int, "text": str}.
    """
    pages = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            cleaned = _clean_text(text)
            if cleaned:
                pages.append({"page": page_num + 1, "text": cleaned})
        doc.close()
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Failed to parse PDF: {e}")
    return pages


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract paragraphs text from a Word document."""
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        # Include table cells if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return _clean_text("\n".join(paragraphs))
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Failed to parse DOCX: {e}")


def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Decode raw TXT bytes to string."""
    try:
        return _clean_text(txt_bytes.decode("utf-8", errors="ignore"))
    except Exception as e:
        logger.error(f"TXT extraction failed: {e}")
        raise ValueError(f"Failed to parse TXT: {e}")


def chunk_document(
    file_bytes: bytes,
    filename: str,
    chunk_size: int = 500,
    chunk_overlap: int = 100,
) -> list[TextChunk]:
    """
    General document chunker.
    Determines type from extension, extracts text, and chunks it using RecursiveCharacterTextSplitter.
    """
    ext = filename.split(".")[-1].lower()
    chunks = []
    chunk_id_counter = 0

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )

    if ext == "pdf":
        pages = extract_text_from_pdf(file_bytes)
        for page_data in pages:
            page_num = page_data["page"]
            split_texts = splitter.split_text(page_data["text"])
            for text in split_texts:
                if len(text.strip()) > 10:
                    chunks.append(
                        TextChunk(
                            text=text,
                            page=page_num,
                            chunk_id=chunk_id_counter,
                        )
                    )
                    chunk_id_counter += 1

    elif ext == "docx":
        text = extract_text_from_docx(file_bytes)
        split_texts = splitter.split_text(text)
        for text in split_texts:
            if len(text.strip()) > 10:
                chunks.append(
                    TextChunk(
                        text=text,
                        page=1,  # Default page for word documents
                        chunk_id=chunk_id_counter,
                    )
                )
                chunk_id_counter += 1

    elif ext == "txt":
        text = extract_text_from_txt(file_bytes)
        split_texts = splitter.split_text(text)
        for text in split_texts:
            if len(text.strip()) > 10:
                chunks.append(
                    TextChunk(
                        text=text,
                        page=1,  # Default page for text files
                        chunk_id=chunk_id_counter,
                    )
                )
                chunk_id_counter += 1
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    logger.info(f"Created {len(chunks)} chunks from document '{filename}'")
    return chunks


def extract_raw_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract all text from a document file and return as a single string.
    No chunking — used for summary generation.
    """
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        pages = extract_text_from_pdf(file_bytes)
        return "\n\n".join(p["text"] for p in pages)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def get_page_count(file_bytes: bytes, filename: str) -> int:
    """Return page count for PDFs, or 1 for other documents."""
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            count = len(doc)
            doc.close()
            return count
        except Exception as e:
            raise ValueError(f"Could not read PDF: {e}")
    return 1
